import os
from docx import Document
from database.models import Category, Test

CATALOG_PATH = "catalog.docx"

async def update_catalog():
    document = Document()
    document.add_heading('Testlar Katalogi', 0)

    categories = await Category.all()
    
    for category in categories:
        document.add_heading(category.name, level=1)
        
        tests = await Test.filter(category_id=category.id).all()
        if not tests:
            document.add_paragraph("Testlar mavjud emas.", style='List Bullet')
            continue
            
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Test Nomi'
        hdr_cells[1].text = 'Kod'
        
        for test in tests:
            row_cells = table.add_row().cells
            row_cells[0].text = test.name
            # Code format: CategoryID-TestID
            row_cells[1].text = f"{category.id}-{test.id}"

    document.save(CATALOG_PATH)
    return CATALOG_PATH
